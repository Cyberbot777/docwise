import os
import json
import fitz
from pathlib import Path
from docx import Document as DocxReader
from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv

load_dotenv()

DATA_DIR = Path(__file__).parent.parent / "data"
FAISS_DIR = Path(__file__).parent.parent / "faiss_index"

def clean_text(text: str) -> str:
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.lower().startswith(("page", "---", "___", "table of contents")):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)

def load_and_embed_documents():
    docs = []

    for file in DATA_DIR.glob("*"):
        ext = file.suffix.lower()

        try:
            if ext in [".txt", ".md"]:
                loader = TextLoader(str(file), encoding="utf-8", autodetect_encoding=True)
                loaded = loader.load()
                text = loaded[0].page_content if loaded else ""
                if len(text.strip()) < 50:
                    print(f"⚠️ Skipping {file.name} — too little content.")
                    continue
                docs.append(Document(page_content=clean_text(text), metadata={"source": file.name}))

            elif ext == ".pdf":
                with fitz.open(file) as pdf:
                    text = "\n".join(page.get_text() for page in pdf)
                    if len(text.strip()) < 50:
                        print(f"⚠️ Skipping {file.name} — too little content.")
                        continue
                    docs.append(Document(page_content=clean_text(text), metadata={"source": file.name}))

            elif ext == ".json":
                with open(file, "r", encoding="utf-8") as f:
                    content = json.load(f)

                def extract_strings(obj):
                    if isinstance(obj, dict):
                        return [v for val in obj.values() for v in extract_strings(val)]
                    elif isinstance(obj, list):
                        return [v for item in obj for v in extract_strings(item)]
                    elif isinstance(obj, str):
                        return [obj]
                    return []

                all_strings = extract_strings(content)
                combined = "\n".join(all_strings)
                if len(combined.strip()) < 50:
                    print(f"⚠️ Skipping {file.name} — too little content.")
                    continue
                docs.append(Document(page_content=clean_text(combined), metadata={"source": file.name}))

            elif ext == ".docx":
                docx = DocxReader(str(file))
                text = "\n".join(p.text for p in docx.paragraphs)
                if len(text.strip()) < 50:
                    print(f"⚠️ Skipping {file.name} — too little content.")
                    continue
                docs.append(Document(page_content=clean_text(text), metadata={"source": file.name}))

            else:
                print(f"⚠️ Unsupported file type: {file.name}")

        except Exception as e:
            print(f"❌ Failed to load {file.name}: {e}")

    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_documents(docs)
    print(f"📄 Embedded {len(chunks)} chunks from {len(docs)} documents.")

    embeddings = OpenAIEmbeddings()

    if FAISS_DIR.exists():
        for f in FAISS_DIR.glob("*"):
            f.unlink()
        FAISS_DIR.rmdir()

    vectorstore = FAISS.from_documents(chunks, embedding=embeddings)
    vectorstore.save_local(str(FAISS_DIR))
    return vectorstore

def ask_question(question: str) -> str:
    embeddings = OpenAIEmbeddings()
    llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0)

    if not FAISS_DIR.exists():
        vectorstore = load_and_embed_documents()
    else:
        vectorstore = FAISS.load_local(
            str(FAISS_DIR),
            embeddings,
            allow_dangerous_deserialization=True
        )

    docs = vectorstore.similarity_search(question, k=6)
    context = "\n\n".join(doc.page_content for doc in docs)

    if not context.strip():
        context = "No relevant documents were found. Please attempt to answer anyway based on your training."

    prompt = f"""
Use the following context to answer the user's question.
If the answer cannot be found, say "I couldn't find the answer in the provided documents."

Context:
{context}

Question:
{question}
"""

    response = llm.invoke(prompt)
    return response.content
